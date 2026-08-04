# =====================================================
# RPM CERDAS AI OFFLINE - EDISI PENYEMPURNAAN MENDALAM
# Generator Rencana Pembelajaran Mendalam
# Kurikulum Merdeka SMP
# =====================================================

import io
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt

# =====================================================
# KONFIGURASI APLIKASI
# =====================================================

APP_NAME = "RPM CERDAS AI"

st.set_page_config(
    page_title=APP_NAME,
    page_icon="📘",
    layout="wide"
)

# =====================================================
# HEADER APLIKASI
# =====================================================

st.markdown(
    """
    <h1 style="text-align:center;">
    📘 RPM CERDAS AI
    </h1>
    <h4 style="text-align:center;">
    Generator Rencana Pembelajaran Mendalam (Offline & Terperinci)
    Kurikulum Merdeka SMP
    </h4>
    """,
    unsafe_allow_html=True
)

st.divider()

# =====================================================
# IDENTITAS GURU
# =====================================================

st.subheader("🏫 Identitas Pembelajaran")

kolom1, kolom2 = st.columns(2)

with kolom1:
    sekolah = st.text_input("Nama Sekolah", value="SMP Negeri")
    guru = st.text_input("Nama Guru")
    mapel = st.selectbox(
        "Mata Pelajaran",
        [
            "Pendidikan Agama",
            "PPKn",
            "Bahasa Indonesia",
            "Matematika",
            "IPA",
            "IPS",
            "Bahasa Inggris",
            "Seni Budaya",
            "PJOK",
            "Informatika",
            "Prakarya"
        ]
    )
    kelas = st.selectbox("Kelas", ["VII", "VIII", "IX"])

with kolom2:
    semester = st.selectbox("Semester", ["Ganjil", "Genap"])
    tahun_pelajaran = st.text_input("Tahun Pelajaran", value="2026/2027")
    alokasi_waktu = st.text_input("Alokasi Waktu", value="2 x 40 Menit")
    model_pembelajaran = st.selectbox(
        "Model Pembelajaran",
        [
            "Problem Based Learning",
            "Project Based Learning",
            "Discovery Learning",
            "Inquiry Learning",
            "Cooperative Learning"
        ]
    )

st.divider()

# =====================================================
# DATA MATERI
# =====================================================

st.subheader("📚 Materi Pembelajaran")

topik = st.text_input("Topik Pembelajaran")
sub_topik = st.text_input("Sub Topik (Opsional)")

cp = st.text_area(
    "Capaian Pembelajaran (CP)",
    height=120,
    placeholder="Masukkan CP Kurikulum Merdeka yang relevan..."
)

karakteristik = st.text_area(
    "Karakteristik Peserta Didik / Kesiapan Belajar (Opsional)",
    height=100,
    placeholder="Contoh: Peserta didik memiliki gaya belajar variatif (visual & kinestetik), sebagian masih memerlukan bimbingan membaca pemahaman..."
)

tujuan_manual = st.text_area(
    "Tujuan Pembelajaran (Opsional, Kosongkan agar dibuat otomatis)",
    height=100,
    placeholder="Masukkan Tujuan Pembelajaran jika Anda ingin menuliskannya secara manual..."
)

st.info(
    """
    💡 Versi ini menggunakan kecerdasan berbasis aturan (rule-based) yang dirancang mendalam.
    Tidak membutuhkan API Key atau koneksi internet.
    """
)

st.divider()

# =====================================================
# ENGINE GENERATOR RPM OFFLINE (TERPERINCI)
# =====================================================

def buat_rpm_offline(
    mapel,
    kelas,
    topik,
    sub_topik,
    cp,
    karakteristik,
    model,
    tujuan_manual
):
    sub_topik_str = f" - Sub Topik: {sub_topik}" if sub_topik.strip() else ""
    cp_clean = cp.strip() if cp.strip() else "Peserta didik memahami konsep materi esensial sesuai standar kurikulum."
    
    # 1. TUJUAN PEMBELAJARAN (ABCD Format & Lebih Detil)
    if tujuan_manual.strip():
        tujuan = tujuan_manual
    else:
        tujuan = f"""Berdasarkan Capaian Pembelajaran (CP):
"{cp_clean}"

Tujuan Pembelajaran (TP) yang dijabarkan secara rinci:
1. Melalui kegiatan eksplorasi informasi dan penyelidikan berkelompok (Condition), peserta didik (Audience) mampu menganalisis konsep dasar mengenai {topik}{sub_topik_str} (Behavior) secara kritis dan mandiri dengan ketepatan minimal 80% (Degree).
2. Melalui penerapan model {model} (Condition), peserta didik (Audience) mampu menyajikan laporan atau produk pemecahan masalah kreatif terkait {topik} (Behavior) secara bergotong royong sesuai dengan kriteria penilaian yang disepakati (Degree).
3. Selama proses pembelajaran (Condition), peserta didik (Audience) dapat menunjukkan sikap bernalar kritis, tanggung jawab, dan saling menghargai (Behavior) secara konsisten (Degree)."""

    # 2. DIMENSI PROFIL PELAJAR PANCASILA
    dimensi = f"""Penguatan Profil Pelajar Pancasila yang ditargetkan pada materi {topik}:

1. Bernalar Kritis: Peserta didik mengidentifikasi gagasan, mengklarifikasi argumen, dan menganalisis informasi yang relevan untuk memecahkan masalah terkait {topik}.
2. Gotong Royong: Peserta didik membangun komunikasi yang efektif, menyelaraskan tindakan kelompok, dan berbagi peran demi menyukseskan tugas bersama.
3. Mandiri: Peserta didik mengambil inisiatif belajar, memantau kemajuan belajarnya sendiri, dan bertanggung jawab atas hasil kerja yang dicapai."""

    # 3. PRAKTIK PEDAGOGIS (Pembelajaran Berdiferensiasi & Sintaks Model)
    karakteristik_str = karakteristik.strip() if karakteristik.strip() else "Heterogen dengan minat dan kesiapan belajar yang bervariasi."
    
    # Strategi Sintaksis Model Pembelajaran
    sintaks_model = ""
    if model == "Problem Based Learning":
        sintaks_model = """Sintaks Problem Based Learning (PBL):
- Tahap 1: Orientasi peserta didik pada masalah kontekstual.
- Tahap 2: Mengorganisasikan peserta didik untuk mendefinisikan tugas belajar.
- Tahap 3: Membimbing penyelidikan individu maupun kelompok kecil.
- Tahap 4: Mengembangkan dan menyajikan hasil karya (solusi/laporan).
- Tahap 5: Menganalisis dan mengevaluasi proses pemecahan masalah."""
    elif model == "Project Based Learning":
        sintaks_model = """Sintaks Project Based Learning (PjBL):
- Tahap 1: Menentukan pertanyaan mendasar atau tantangan produk.
- Tahap 2: Merancang perencanaan langkah-langkah pembuatan proyek.
- Tahap 3: Menyusun jadwal pengerjaan proyek (timeline).
- Tahap 4: Memonitor keaktifan dan perkembangan draf proyek siswa.
- Tahap 5: Menguji hasil (presentasi produk/karya).
- Tahap 6: Mengevaluasi pengalaman belajar dan memberikan umpan balik."""
    elif model == "Discovery Learning":
        sintaks_model = """Sintaks Discovery Learning:
- Tahap 1: Stimulation (Pemberian rangsangan/fenomena menarik).
- Tahap 2: Problem Statement (Identifikasi masalah dan penyusunan hipotesis).
- Tahap 3: Data Collection (Pengumpulan data dan literatur mandiri).
- Tahap 4: Data Processing (Pengolahan data dan klasifikasi informasi).
- Tahap 5: Verification (Pembuktian kesesuaian data dengan teori).
- Tahap 6: Generalization (Menarik kesimpulan umum)."""
    elif model == "Inquiry Learning":
        sintaks_model = """Sintaks Inquiry Learning:
- Tahap 1: Orientasi masalah dan penjelasan batas area penyelidikan.
- Tahap 2: Merumuskan masalah berupa pertanyaan ilmiah.
- Tahap 3: Mengajukan dugaan sementara atau hipotesis awal.
- Tahap 4: Mengumpulkan data pendukung melalui eksperimen/penelusuran literatur.
- Tahap 5: Menguji hipotesis secara analitis.
- Tahap 6: Merumuskan kesimpulan akhir."""
    else:  # Cooperative Learning
        sintaks_model = """Sintaks Cooperative Learning:
- Tahap 1: Menyampaikan tujuan instruksional dan membangkitkan motivasi siswa.
- Tahap 2: Menyajikan materi/informasi awal secara interaktif.
- Tahap 3: Mengorganisasikan siswa ke dalam kelompok belajar heterogen.
- Tahap 4: Membimbing jalannya diskusi kelompok (kolaborasi peran).
- Tahap 5: Melakukan evaluasi atau kuis pemahaman materi.
- Tahap 6: Memberikan apresiasi atau penghargaan atas performa kelompok."""

    pedagogis = f"""A. Pendekatan Utama: Student-Centered Learning & Integrasi TPACK.
B. Karakteristik Peserta Didik: {karakteristik_str}

C. Implementasi Pembelajaran Berdiferensiasi:
1. Diferensiasi Konten: Menyediakan bahan ajar bervariasi (artikel bacaan ringkas, infografis visual, serta video penjelasan materi {topik}).
2. Diferensiasi Proses: Memberikan bimbingan bertingkat (scaffolding) untuk kelompok yang membutuhkan bantuan ekstra, serta menyediakan bahan perluasan bagi siswa berkinerja cepat.
3. Diferensiasi Produk: Mengizinkan peserta didik melaporkan hasil kerja kelompok dalam bentuk draf laporan tertulis, mind map kreatif, atau infografis digital sesuai minat kelompok.

D. Langkah Strategis Model Pembelajaran:
{sintaks_model}"""

    # 4. LINGKUNGAN PEMBELAJARAN
    lingkungan = f"""Penataan & Pengkondisian Kelas Fisik dan Non-Fisik:

1. Tata Letak Fleksibel: Meja dan kursi diatur dalam bentuk kelompok kooperatif (lingkaran kecil/berhadapan) guna memudahkan kolaborasi aktif tanpa membatasi ruang gerak guru sebagai fasilitator.
2. Iklim Psikologis Aman & Inklusif: Membangun kesepakatan kelas yang mengedepankan apresiasi, toleransi pendapat, dan peniadaan diskriminasi, sehingga siswa berani mencoba tanpa takut melakukan kesalahan.
3. Pojok Informasi: Memanfaatkan papan tulis atau dinding kelas sebagai area pajang draf kerja siswa sebagai bentuk apresiasi langsung."""

    # 5. KEMITRAAN PEMBELAJARAN
    kemitraan = f"""Bentuk Kemitraan yang Diimplementasikan:

1. Kemitraan dengan Orang Tua: Mengirimkan info penugasan kelompok {topik} melalui media komunikasi agar orang tua dapat memantau atau mengarahkan aktivitas anak di rumah.
2. Tutor Sebaya (Peer-Teaching): Menunjuk siswa yang lebih cepat memahami konsep untuk membantu mendampingi rekan sekelompoknya dalam proses penemuan informasi.
3. Kemitraan Lingkungan Sekolah: Melibatkan perpustakaan sekolah atau area taman (jika relevan) sebagai sarana eksplorasi materi di luar kelas fisik."""

    # 6. PEMANFAATAN DIGITAL (Sesuai Karakteristik Mata Pelajaran)
    contoh_media = {
        "Pendidikan Agama": "Aplikasi Al-Qur'an/Kitab Digital, Canva, Video Keteladanan.",
        "PPKn": "Situs berita resmi, Padlet untuk opini publik, YouTube.",
        "Bahasa Indonesia": "Google Docs untuk kolaborasi menulis teks, Canva, KBBI daring.",
        "Matematika": "GeoGebra, Desmos Calculator, Mentimeter.",
        "IPA": "Virtual Lab (PhET Simulation), video eksperimen, Google Lens.",
        "IPS": "Google Earth, Atlas digital, infografis peristiwa bersejarah.",
        "Bahasa Inggris": "Lirik lagu interaktif, Google Translate (analisis semantik), video percakapan nyata.",
        "Seni Budaya": "Platform galeri virtual museum, Canva desain, Pinterest.",
        "PJOK": "Aplikasi perekam video lambat (slow-motion), pelacak kebugaran sederhana.",
        "Informatika": "Platform Scratch, Replit, emulator jaringan, Google Colab.",
        "Prakarya": "Platform Pinterest, video tutorial DIY, Canva kemasan produk."
    }
    media_spesifik = contoh_media.get(mapel, "Google Slides, Canva, Quizizz, Video Pembelajaran.")

    digital = f"""Integrasi Teknologi Digital (TPACK) dalam Pembelajaran {mapel}:

1. Akses Informasi Mandiri: Memanfaatkan gawai siswa untuk mengakses bahan ajar digital atau e-book materi {topik}.
2. Media Kolaborasi: Menggunakan media presentasi interaktif seperti Canva Pendidikan untuk menyusun laporan/produk akhir.
3. Media Interaktif & Evaluasi: Menggunakan platform Quizizz atau Kahoot untuk asesmen formatif akhir yang menyenangkan.
4. Alat Spesifik Pembelajaran: {media_spesifik}"""

    # 7. LANGKAH PEMBELAJARAN DETAIL (Dengan rincian waktu konkret)
    langkah_inti = ""
    if model == "Problem Based Learning":
        langkah_inti = f"""- Tahap 1: Orientasi Masalah (15 Menit)
  Guru menayangkan kasus/video faktual terkait permasalahan riil {topik}. Peserta didik merumuskan pertanyaan penting mengenai masalah tersebut.
- Tahap 2: Mengorganisasikan Siswa (10 Menit)
  Peserta didik membentuk kelompok heterogen. Guru membagikan LKPD (Lembar Kerja Peserta Didik) dan membagi peran kerja tim.
- Tahap 3: Penyelidikan Terbimbing (20 Menit)
  Peserta didik mengumpulkan informasi dari buku paket dan internet. Guru melakukan scaffolding (bimbingan terarah) bagi kelompok yang kesulitan.
- Tahap 4: Menyusun & Menyajikan Karya (15 Menit)
  Siswa berkolaborasi menyusun hasil pemecahan masalah dalam bentuk bahan presentasi (infografis/mind map).
- Tahap 5: Analisis & Evaluasi Solusi (10 Menit)
  Perwakilan kelompok mempresentasikan hasil. Kelompok lain memberikan masukan konstruktif. Guru memberikan konfirmasi keilmuan."""
    elif model == "Project Based Learning":
        langkah_inti = f"""- Tahap 1: Pertanyaan Mendasar (10 Menit)
  Guru memberikan tantangan nyata yang membutuhkan produk pemecahan masalah berkaitan dengan {topik}.
- Tahap 2: Mendesain Perencanaan Proyek (15 Menit)
  Siswa di dalam kelompok menyusun ide proyek, menetapkan bahan-bahan, serta membagi tugas kerja masing-masing.
- Tahap 3: Menyusun Jadwal (10 Menit)
  Kelompok bersama guru membuat jadwal timeline penyelesaian (mencakup batas draf awal hingga penyelesaian akhir).
- Tahap 4: Memonitor Progres Proyek (15 Menit)
  Siswa mulai merakit draf/sketsa awal produk, guru berkeliling memeriksa kendala teknis kelompok.
- Tahap 5: Menguji Hasil / Presentasi (15 Menit)
  Setiap kelompok memaparkan produk awal/hasil rancangan di depan kelas untuk menerima umpan balik.
- Tahap 6: Evaluasi Proses Belajar (10 Menit)
  Siswa dan guru merefleksikan keberhasilan serta hambatan teknis pembuatan proyek."""
    elif model == "Discovery Learning":
        langkah_inti = f"""- Tahap 1: Stimulasi/Rangsangan (10 Menit)
  Guru memajang gambar/konsep yang memicu kontradiksi berpikir terkait {topik}.
- Tahap 2: Identifikasi Masalah (10 Menit)
  Siswa mengidentifikasi pertanyaan kunci dan merumuskan hipotesis jawaban sementara.
- Tahap 3: Pengumpulan Data (20 Menit)
  Siswa secara berkelompok mencari informasi dari berbagai sumber (buku, bahan ajar, internet).
- Tahap 4: Pengolahan Data (15 Menit)
  Siswa mengklasifikasikan temuan, mendiskusikan korelasi data, dan menyusunnya secara sistematis di LKPD.
- Tahap 5: Verifikasi/Pembuktian (15 Menit)
  Siswa membuktikan kebenaran hipotesis mereka dengan membandingkannya terhadap kajian literatur teoretis.
- Tahap 6: Generalisasi (10 Menit)
  Siswa menyusun kesimpulan akhir bersama mengenai konsep {topik}."""
    elif model == "Inquiry Learning":
        langkah_inti = f"""- Tahap 1: Orientasi Fokus (10 Menit)
  Guru menyampaikan arahan ruang lingkup observasi serta tujuan penemuan konsep {topik}.
- Tahap 2: Merumuskan Masalah (10 Menit)
  Siswa membuat rumusan masalah terarah mengenai aspek yang diselidiki.
- Tahap 3: Pengajuan Hipotesis (10 Menit)
  Kelompok menyusun dugaan logis awal berdasarkan pengetahuan awal yang mereka miliki.
- Tahap 4: Pengumpulan Data Lapangan/Studi (20 Menit)
  Siswa melangsungkan penyelidikan mandiri, melakukan eksperimen, atau penelusuran pustaka intensif.
- Tahap 5: Menguji Hipotesis (15 Menit)
  Kelompok memverifikasi kesesuaian data yang mereka temukan di lapangan dengan hipotesis awal.
- Tahap 6: Kesimpulan & Laporan (15 Menit)
  Siswa memformulasikan kesimpulan, menyusun laporan ringkas, dan membagikannya ke hadapan kelas."""
    else:  # Cooperative Learning
        langkah_inti = f"""- Tahap 1: Penjelasan Tujuan & Motivasi (10 Menit)
  Guru menyampaikan skema belajar kooperatif hari ini serta pentingnya kolaborasi aktif.
- Tahap 2: Penyajian Informasi (15 Menit)
  Guru memberikan penjelasan singkat materi {topik} agar siswa mendapatkan bekal konsep awal.
- Tahap 3: Pembagian Kelompok (10 Menit)
  Siswa berkumpul dengan kelompok heterogen yang telah dibagikan peran masing-masing anggota.
- Tahap 4: Aktivitas Kelompok (25 Menit)
  Siswa bekerja sama menyelesaikan LKPD kelompok (contoh metode: Jigsaw / STAD). Guru mendampingi jalannya proses diskusi.
- Tahap 5: Evaluasi Bersama (15 Menit)
  Perwakilan siswa memaparkan jawaban, kelompok lain saling bertukar koreksi. Guru memberikan kuis formatif singkat.
- Tahap 6: Pemberian Apresiasi (5 Menit)
  Guru memberikan penghargaan atas kinerja kerja sama kelompok terbaik."""

    langkah = f"""A. PENDAHULUAN (15 Menit)

1. Persiapan Kelas:
   - Guru membuka kelas dengan salam ramah, menanyakan kondisi siswa, berdoa bersama dipimpin ketua kelas, dan memeriksa kerapian kelas serta kehadiran siswa.
2. Apersepsi:
   - Guru mengaitkan konsep materi yang akan dipelajari dengan materi di pertemuan sebelumnya atau pengalaman sehari-hari siswa.
   - Mengajukan pertanyaan pemantik kontekstual: "Pernahkah kalian menemui situasi di mana...?" guna menarik fokus perhatian siswa.
3. Orientasi Tujuan:
   - Menyampaikan target kompetensi (ATP) dan memberikan penjelasan singkat mengenai skema penilaian proses hari ini.

B. KEGIATAN INTI (~50 - 70 Menit) - Penerapan Model {model}

{langkah_inti}

C. PENUTUP (15 Menit)

1. Simpulan Bersama:
   - Guru membimbing siswa secara kolaboratif menyimpulkan konsep utama yang telah dikonstruksi hari ini mengenai materi {topik}.
2. Refleksi Pembelajaran:
   - Refleksi siswa: Mengajukan pertanyaan reflektif seperti "Bagian mana dari pelajaran hari ini yang paling menantang?" dan "Bagaimana cara kerja sama kelompok kita hari ini?"
   - Refleksi guru: Evaluasi singkat efektivitas penyampaian materi dan model pembelajaran.
3. Rencana Tindak Lanjut:
   - Menginformasikan materi atau persiapan untuk pertemuan mendatang.
   - Kelas diakhiri dengan doa bersama serta salam penutup."""

    # 8. ASESMEN PEMBELAJARAN (Sangat Spesifik & Terstruktur)
    asesmen_spesifik = {
        "Matematika": "Tes Formatif tertulis (Analisis masalah kontekstual), Lembar penilaian kinerja pemecahan soal terstruktur.",
        "IPA": "Laporan praktikum ilmiah, Rubrik presentasi hasil investigasi kelompok, Tes tertulis HOTS.",
        "Bahasa Indonesia": "Asesmen produk tulisan/teks, Rubrik penilaian unjuk kerja membaca/presentasi.",
        "Bahasa Inggris": "Rubrik performa berbicara (speaking performance), Lembar analisis teks bacaan.",
        "Informatika": "Demonstrasi portofolio kode/desain sistem digital kelompok, Tes formatif praktik.",
        "IPS": "Asesmen produk infografis/peta konsep historis, Analisis studi kasus isu sosial.",
        "PPKn": "Lembar penilaian proyek kampanye kewarganegaraan, Portofolio telaah kasus hak dan kewajiban.",
        "PJOK": "Lembar pengamatan unjuk kerja keterampilan gerak, Penilaian diri aspek kebugaran jasmani."
    }
    asesmen_mapel = asesmen_spesifik.get(mapel, "Tes tertulis pemahaman konsep, Rubrik presentasi produk akhir kelompok.")

    asesmen = f"""A. ASESMEN DIAGNOSTIK (Awal)
1. Non-Kognitif: Kuesioner singkat atau observasi lisan di awal pelajaran guna memetakan suasana hati siswa dan kesiapan mental belajar hari ini.
2. Kognitif: Memberikan 2-3 soal singkat/pertanyaan pemantik terkait materi prasyarat sebelum memasuki bahasan utama {topik}.

B. ASESMEN FORMATIF (Selama Proses Pembelajaran)
1. Asesmen Perilaku (Sikap):
   * Menggunakan Lembar Observasi Sikap Profil Pelajar Pancasila untuk menilai keaktifan Bernalar Kritis, Gotong Royong, dan Kemandirian selama diskusi kelompok.
2. Asesmen Kinerja Proses (Diskusi & Presentasi):
   * Rubrik Penilaian Kinerja Kelompok:
     - Kerjasama Tim (Skor 1 - 4)
     - Pemahaman Materi {topik} (Skor 1 - 4)
     - Kemampuan Komunikasi/Presentasi (Skor 1 - 4)

C. ASESMEN SUMATIF (Hasil Akhir)
1. Jenis Penilaian: {asesmen_mapel}
2. Rubrik Penilaian Produk / Laporan Akhir:
   - Kesesuaian Konsep dengan Capaian Pembelajaran (Bobot 40%)
   - Kreativitas dan Estetika Sajian/Media (Bobot 30%)
   - Struktur dan Kejelasan Penyusunan (Bobot 30%)

D. PROGRAM TINDAK LANJUT
1. Pembelajaran Pengayaan: Diberikan kepada peserta didik yang telah melampaui kriteria ketuntasan, berupa tugas analisis mendalam studi kasus pemecahan masalah (HOTS).
2. Pembelajaran Remedial: Diberikan kepada peserta didik yang belum tuntas, berupa bimbingan terarah (scaffolding) individual atau pemanfaatan tutor sebaya pada submateri yang belum dipahami."""

    return {
        "dimensi_profil": dimensi,
        "tujuan_pembelajaran": tujuan,
        "praktik_pedagogis": pedagogis,
        "lingkungan_belajar": lingkungan,
        "kemitraan_belajar": kemitraan,
        "pemanfaatan_digital": digital,
        "langkah_pembelajaran": langkah,
        "asesmen_total": asesmen
    }

# =====================================================
# TOMBOL GENERATE RPM
# =====================================================

if st.button(
    "🚀 BUAT RPM CERDAS AI OFFLINE",
    type="primary",
    use_container_width=True
):
    if not topik.strip():
        st.warning("Silakan isi Topik Pembelajaran terlebih dahulu.")
        st.stop()

    hasil_rpm = buat_rpm_offline(
        mapel=mapel,
        kelas=kelas,
        topik=topik,
        sub_topik=sub_topik,
        cp=cp,
        karakteristik=karakteristik,
        model=model_pembelajaran,
        tujuan_manual=tujuan_manual
    )

    st.session_state.hasil_rpm = hasil_rpm
    st.success("✅ RPM berhasil dibuat secara terperinci.")

# =====================================================
# MENAMPILKAN HASIL RPM
# =====================================================

if "hasil_rpm" in st.session_state:
    st.divider()
    st.subheader("📄 HASIL RPM CERDAS AI OFFLINE")

    hasil = st.session_state.hasil_rpm

    daftar_komponen = [
        ("1. Dimensi Profil Lulusan", "dimensi_profil"),
        ("2. Tujuan Pembelajaran", "tujuan_pembelajaran"),
        ("3. Praktik Pedagogis", "praktik_pedagogis"),
        ("4. Lingkungan Pembelajaran", "lingkungan_belajar"),
        ("5. Kemitraan Pembelajaran", "kemitraan_belajar"),
        ("6. Pemanfaatan Digital", "pemanfaatan_digital"),
        ("7. Langkah Pembelajaran", "langkah_pembelajaran"),
        ("8. Asesmen Pembelajaran", "asesmen_total")
    ]

    for judul, kode in daftar_komponen:
        with st.expander(judul, expanded=False):
            st.write(hasil.get(kode, ""))

    st.divider()

    st.subheader("📌 Ringkasan Identitas RPM")

    tabel_identitas = {
        "Nama Sekolah": sekolah,
        "Nama Guru": guru,
        "Mata Pelajaran": mapel,
        "Kelas": kelas,
        "Semester": semester,
        "Tahun Pelajaran": tahun_pelajaran,
        "Alokasi Waktu": alokasi_waktu,
        "Topik": topik,
        "Sub Topik": sub_topik if sub_topik.strip() else "-"
    }

    for k, v in tabel_identitas.items():
        st.write(f"**{k}:** {v}")

# =====================================================
# FUNGSI MEMBUAT DOKUMEN WORD
# =====================================================

def buat_dokumen_rpm(data):
    doc = Document()

    # PENGATURAN HALAMAN
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # JUDUL
    judul = doc.add_heading(level=0)
    judul.alignment = 1
    run = judul.add_run("RENCANA PEMBELAJARAN MENDALAM (RPM)")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # IDENTITAS
    doc.add_heading("I. IDENTITAS PEMBELAJARAN", level=2)

    identitas = [
        ("Nama Sekolah", data["sekolah"]),
        ("Nama Guru", data["guru"]),
        ("Mata Pelajaran", data["mapel"]),
        ("Kelas / Semester", f'{data["kelas"]} / {data["semester"]}'),
        ("Tahun Pelajaran", data["tahun_pelajaran"]),
        ("Alokasi Waktu", data["alokasi_waktu"]),
        ("Topik", data["topik"])
    ]

    tabel = doc.add_table(rows=len(identitas), cols=2)
    tabel.style = "Table Grid"

    for i, (a, b) in enumerate(identitas):
        tabel.rows[i].cells[0].text = a
        tabel.rows[i].cells[1].text = str(b)

    doc.add_paragraph()

    # KOMPONEN RPM
    doc.add_heading("II. KOMPONEN RPM", level=2)

    komponen = [
        ("1. Dimensi Profil Lulusan", data["dimensi_profil"]),
        ("2. Tujuan Pembelajaran", data["tujuan_pembelajaran"]),
        ("3. Praktik Pedagogis", data["praktik_pedagogis"]),
        ("4. Lingkungan Pembelajaran", data["lingkungan_belajar"]),
        ("5. Kemitraan Pembelajaran", data["kemitraan_belajar"]),
        ("6. Pemanfaatan Digital", data["pemanfaatan_digital"]),
        ("7. Langkah Pembelajaran", data["langkah_pembelajaran"]),
        ("8. Asesmen Pembelajaran", data["asesmen_total"])
    ]

    for judul_sub, isi in komponen:
        doc.add_heading(judul_sub, level=3)
        doc.add_paragraph(str(isi))

    # PENGESAHAN
    doc.add_heading("III. PENGESAHAN", level=2)
    tanda_tangan = doc.add_table(rows=1, cols=2)
    tanda_tangan.style = "Table Grid"

    tanda_tangan.cell(0, 0).text = (
        "Mengetahui,\n\n"
        "Kepala Sekolah\n\n\n\n"
        "(............................)"
    )

    tanda_tangan.cell(0, 1).text = (
        "Guru Mata Pelajaran\n\n\n\n"
        f"({data['guru']})"
    )

    file = io.BytesIO()
    doc.save(file)
    file.seek(0)

    return file

# =====================================================
# DOWNLOAD WORD
# =====================================================

if "hasil_rpm" in st.session_state:
    hasil = st.session_state.hasil_rpm

    data_word = {
        "sekolah": sekolah,
        "guru": guru,
        "mapel": mapel,
        "kelas": kelas,
        "semester": semester,
        "tahun_pelajaran": tahun_pelajaran,
        "alokasi_waktu": alokasi_waktu,
        "topik": topik,
        **hasil
    }

    file_word = buat_dokumen_rpm(data_word)

    st.download_button(
        label="⬇️ DOWNLOAD RPM WORD (.DOCX)",
        data=file_word,
        file_name="RPM_CERDAS_AI_MENDALAM.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# =====================================================
# TOMBOL RESET
# =====================================================

st.divider()

if st.button("🔄 HAPUS HASIL RPM", use_container_width=True):
    if "hasil_rpm" in st.session_state:
        del st.session_state["hasil_rpm"]

    st.success("Hasil RPM berhasil dihapus.")
    st.rerun()

# =====================================================
# FOOTER
# =====================================================

st.divider()

st.markdown(
    """
    <center>
    📘 <b>RPM CERDAS AI OFFLINE</b><br>
    Generator Rencana Pembelajaran Mendalam<br>
    Kurikulum Merdeka SMP
    </center>
    """,
    unsafe_allow_html=True
)
